from PyQt5.QtWidgets import (QWidget, QSpinBox, QCheckBox, QTableWidget, QPushButton, QLabel,
                            QDialog, QMessageBox, QTableWidgetItem, QLineEdit, QComboBox, QTextBrowser)
from PyQt5 import uic
import os
from docx.shared import Pt
from docxtpl import DocxTemplate
import re
import json
import sys
from PyQt5.QtWidgets import QFileDialog, QMessageBox
from num2words import num2words
from jinja2 import Environment, DebugUndefined, Template
import html
from docx import Document
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from ui.payment_dialog import PaymaentDialog
from logic.license_checker import is_license_valid
from ui.method__dialog import MethodRejectionDialog

from PyQt5.QtGui import QFont
from ui.cost_method_dialogs.building_choose import BuildingChooseDialog
from ui.cost_method_dialogs.deviations_and_wear_dialog import DeviationsAndWearDialog

from logic.data_entry import DataEntryForm
from PyQt5.QtCore import QDate, Qt
import json
import pandas as pd
from functools import partial
from logic.paths import get_ui_path

import traceback

class AgreementWidget(QWidget):
    def __init__(self, parent=None, main_window=None, valuation_window=None):
        super().__init__(parent)
        self.main_window = main_window
        self.valuation_window = valuation_window
        uic.loadUi(get_ui_path("agreement_widget.ui"), self)

        
        self.data_service = DataEntryForm()

        self.comboBox_agreement_list = self.findChild(QComboBox, 'comboBox_agreement_list')
        self.checkBox_cost = self.findChild(QCheckBox, 'checkBox_cost')
        # self.checkBox_income = self.findChild(QCheckBox, 'checkBox_income')
        self.checkBox_comparative = self.findChild(QCheckBox, 'checkBox_comparative')
        self.spinBox_cost_percent = self.findChild(QSpinBox, 'spinBox_cost_percent')
        # self.spinBox_income_percent = self.findChild(QSpinBox, 'spinBox_income_percent')
        self.spinBox_comparative_percent = self.findChild(QSpinBox, 'spinBox_comparative_percent')
        self.label_cost_weighted_average = self.findChild(QLabel, 'label_cost_weighted_average')
        # self.label_income_weighted_average = self.findChild(QLabel, 'label_income_weighted_average')
        self.label_comparative_weighted_average = self.findChild(QLabel, 'label_comparative_weighted_average')
        self.label_final_cost = self.findChild(QLabel, 'label_final_cost')
        self.label_building_land = self.findChild(QLabel, 'label_building_land')
        self.pushButton_final_save = self.findChild(QPushButton, 'pushButton_final_save')
        self.pushButton_final_save.clicked.connect(self.final_save)
       
        self.pushButton_upload_report = self.findChild(QPushButton, 'pushButton_upload_report')
        self.pushButton_upload_report.clicked.connect(self.on_upload_report_clicked)

        self.comboBox_agreement_list.addItems(['Взвешенное среднее', 'Среднее арифметическое'])
        self.comboBox_agreement_list.setCurrentText('Взвешенное среднее')
        self.update_agreement_ui()

        self.checkBox_cost.setChecked(False)
        # self.checkBox_income.setChecked(True)
        self.checkBox_comparative.setChecked(False)
        self.checkBox_cost.stateChanged.connect(self.on_checkbox_state_changed)
        #self.checkBox_income.stateChanged.connect(self.on_checkbox_state_changed)
        self.checkBox_comparative.stateChanged.connect(self.on_checkbox_state_changed)

        # Установим начальные значения процентов
        self.spinBox_cost_percent.setValue(50)
        #self.spinBox_income_percent.setValue(33)
        self.spinBox_comparative_percent.setValue(50)

        self.checkBox_cost.stateChanged.connect(self.update_agreement_ui)
        #self.checkBox_income.stateChanged.connect(self.update_agreement_ui)
        self.checkBox_comparative.stateChanged.connect(self.update_agreement_ui)
        self.comboBox_agreement_list.currentTextChanged.connect(self.update_agreement_ui)

        self.spinBox_cost_percent.valueChanged.connect(lambda: self.redistribute_percent("cost"))
        # self.spinBox_income_percent.valueChanged.connect(lambda: self.redistribute_percent("income"))
        self.spinBox_comparative_percent.valueChanged.connect(lambda: self.redistribute_percent("comparative"))
        self.cost_value = 0
        #self.income_value = 0
        self.comparative_value = 0
        self.weighted_cost = 0
        self.weighted_comparative = 0
       

    def format_sum(self, value):
        return f"{round(value):,}".replace(",", " ")

   


    def load_costs_from_json(self, full_data):
        try:
            # Затратный подход
            liters = full_data.get("liters", [])
            building_cost = sum(liter.get("final_cost", 0) for liter in liters)

            land_text = full_data.get("land_valuation", {}).get("land_total_cost", "")
            land_cost = 0
            if land_text:
                land_cost = float(land_text.split(":")[-1].replace("сум", "").replace(" ", "").replace(",", ""))

            self.cost_value = building_cost + land_cost
            self.checkBox_cost.setText(
                f"Затратный подход: {self.format_sum(self.cost_value)} сум\n")
            self.label_building_land.setText(
                f"Стоимость улучшений: {self.format_sum(building_cost)} сум\n"
                f"Права на землю: {self.format_sum(land_cost)} сум"
            )

            # Сравнительный подход
            comp_text = full_data.get("comparative", {}).get("label_comparative_final_cost", "")
            comp_number = "".join(c for c in comp_text if c.isdigit() or c in ",.")
            self.comparative_value = float(comp_number.replace(" ", "").replace(",", "")) if comp_number else 0
            self.checkBox_comparative.setText(f"Сравнительный подход: {self.format_sum(self.comparative_value)} сум")

            self.building_cost = building_cost
            self.land_cost = land_cost
            self.cost_value = building_cost + land_cost

            # ❌ НЕ вызывать здесь:
            # self.update_agreement_ui()
            # self.on_checkbox_state_changed()

        except Exception as e:
            print("[ERROR] Ошибка загрузки стоимостей из JSON:", e)




    def update_agreement_ui(self):
        method = self.comboBox_agreement_list.currentText()
        use_cost = self.checkBox_cost.isChecked()
        #use_income = self.checkBox_income.isChecked()
        use_comparative = self.checkBox_comparative.isChecked()

        widgets = [
            (self.spinBox_cost_percent, self.label_cost_weighted_average, use_cost),
            (self.spinBox_comparative_percent, self.label_comparative_weighted_average, use_comparative),
        ]


        if sum([use_cost, use_comparative]) <= 1:
            for spin, label, _ in widgets:
                spin.setVisible(False)
                label.setVisible(False)
            self.label_final_cost.setText(self.get_single_cost(use_cost, use_comparative))
            return

        if method == 'Среднее арифметическое':
            for spin, label, active in widgets:
                spin.setVisible(False)
                label.setVisible(active)  # <-- Показываем label'ы только для активных подходов

            self.recalculate_average()

            # Отображаем 50/50 доли (визуально, не по логике расчёта — она уже корректна)
            if self.checkBox_cost.isChecked() and self.cost_value is not None:
                cost_portion = self.cost_value / 2
                self.label_cost_weighted_average.setText(f"{self.format_sum(cost_portion)} сум")

            if self.checkBox_comparative.isChecked() and self.comparative_value is not None:
                comp_portion = self.comparative_value / 2
                self.label_comparative_weighted_average.setText(f"{self.format_sum(comp_portion)} сум")

        else:
            for spin, label, active in widgets:
                spin.setVisible(active)
                label.setVisible(active)
            self.recalculate_weighted_average()

    def get_single_cost(self, use_cost, use_comparative):
        if use_cost and self.cost_value:
            return f"Итоговая стоимость: {self.format_sum(self.cost_value)} сум"
        # if use_income and self.income_value:
        #     return f"Итоговая стоимость: {self.format_sum(self.income_value)} сум"
        if use_comparative and self.comparative_value:
            return f"Итоговая стоимость: {self.format_sum(self.comparative_value)} сум"
        return "Итоговая стоимость: н/д"

    def recalculate_average(self):
        total = 0.0
        count = 0

        if self.checkBox_cost.isChecked() and self.cost_value is not None:
            total += self.cost_value
            count += 1

        if self.checkBox_comparative.isChecked() and self.comparative_value is not None:
            total += self.comparative_value
            count += 1

        if count > 0:
            avg = total / count
            self.label_final_cost.setText(f"Итоговая стоимость: {self.format_sum(avg)} сум")
        else:
            self.label_final_cost.setText("Итоговая стоимость: н/д")


    def recalculate_weighted_average(self):
        parts = []
        total_percent = 0

        if self.checkBox_cost.isChecked() and self.cost_value is not None:
            p = self.spinBox_cost_percent.value()
            parts.append((self.cost_value, p, 'cost'))
            total_percent += p
        # if self.checkBox_income.isChecked() and self.income_value is not None:
        #     p = self.spinBox_income_percent.value()
        #     parts.append((self.income_value, p, 'income'))
        #     total_percent += p
        if self.checkBox_comparative.isChecked() and self.comparative_value is not None:
            p = self.spinBox_comparative_percent.value()
            parts.append((self.comparative_value, p, 'comparative'))
            total_percent += p

        if total_percent != 100:
            self.label_final_cost.setText("Ошибка: сумма процентов ≠ 100%")
            return

        weighted_total = sum(value * percent / 100 for value, percent, _ in parts)
        self.label_final_cost.setText(f"Итоговая стоимость: {self.format_sum(weighted_total)} сум")

        # Обновляем подписи с весовой стоимостью
        for value, percent, key in parts:
            portion = value * percent / 100

            if key == 'cost':
                self.weighted_cost = portion
                self.label_cost_weighted_average.setText(f"{self.format_sum(portion)} сум")
            elif key == 'comparative':
                self.weighted_comparative = portion
                self.label_comparative_weighted_average.setText(f"{self.format_sum(portion)} сум")



    def redistribute_percent(self, changed):
        if self.comboBox_agreement_list.currentText() != 'Взвешенное среднее':
            return

        active_boxes = {
            "cost": self.spinBox_cost_percent if self.checkBox_cost.isChecked() else None,
            "comparative": self.spinBox_comparative_percent if self.checkBox_comparative.isChecked() else None
        }

        active_boxes = {k: box for k, box in active_boxes.items() if box is not None}
        if len(active_boxes) != 2:
            return

        changed_box = active_boxes.get(changed)
        if not changed_box:
            return

        changed_value = changed_box.value()
        other_key = [k for k in active_boxes if k != changed][0]
        other_box = active_boxes[other_key]

        new_other_value = 100 - changed_value
        other_box.blockSignals(True)
        other_box.setValue(new_other_value)
        other_box.blockSignals(False)
        if getattr(self, "_loading", False):
            return

        self.recalculate_weighted_average()



    


    def on_checkbox_state_changed(self):
        if self.comboBox_agreement_list.currentText() != 'Взвешенное среднее':
            return

        active_boxes = {
            "cost": self.spinBox_cost_percent if self.checkBox_cost.isChecked() else None,
            "comparative": self.spinBox_comparative_percent if self.checkBox_comparative.isChecked() else None
        }

        active_boxes = {k: box for k, box in active_boxes.items() if box is not None}
        count = len(active_boxes)

        if count == 2:
            for box in active_boxes.values():
                box.blockSignals(True)
                box.setValue(50)
                box.blockSignals(False)
            self.recalculate_weighted_average()

        elif count == 1:
            only_box = list(active_boxes.values())[0]
            only_box.blockSignals(True)
            only_box.setValue(100)
            only_box.blockSignals(False)
            self.recalculate_weighted_average()
        if getattr(self, "_loading", False):
            return



    def collect_agreement_data(self):
        

        final_cost_text = self.label_final_cost.text()
        clean_text = final_cost_text.replace('\xa0', ' ').replace('\u202f', ' ')  # убираем неразрывные пробелы

        match = re.search(r'([\d\s]+)\s*сум', clean_text.lower())
        edited_cost = match.group(1).strip() if match else ""

        try:
            words = num2words(int(edited_cost.replace(" ", "")), lang='ru').capitalize() + " сум"
        except Exception:
            words = ""

        method = self.comboBox_agreement_list.currentText()
        use_cost = self.checkBox_cost.isChecked()
        use_comparative = self.checkBox_comparative.isChecked()

        if sum([use_cost, use_comparative]) == 1:
            agreement_method_summary = (
                "Оценка рыночной стоимости объекта произведена на основе одного подхода, "
                "который признан наиболее достоверным и обоснованным в условиях доступности информации "
                "и специфики объекта оценки. Согласование результатов не производилось, "
                "поскольку альтернативные подходы не использовались или были признаны нецелесообразными."
            )
        elif method == "Взвешенное среднее":
            agreement_method_summary = (
                "Согласование результатов оценки осуществлено методом взвешенного среднего. "
                "Каждому из использованных подходов был присвоен удельный вес, отражающий степень его применимости "
                "и достоверности в контексте объекта оценки. Итоговая величина стоимости определена как сумма произведений "
                "значений стоимости по каждому подходу и соответствующих весовых коэффициентов."
            )
        else:
            agreement_method_summary = (
                "Согласование результатов оценки выполнено методом простого среднего, при котором каждый из выбранных "
                "подходов имеет равную значимость. Итоговая стоимость определена как среднее арифметическое значений, "
                "полученных по отдельным подходам, что обеспечивает сбалансированную и объективную оценку с учётом всех доступных данных."
            )

        # Получаем стоимости из атрибутов
        building_cost_raw = getattr(self, "building_cost", 0)
        building_cost = self.format_sum(building_cost_raw)

        land_cost_raw = getattr(self, "land_cost", 0)
        land_cost = self.format_sum(land_cost_raw)

        total_cost_value_raw = getattr(self, "cost_value", 0)
        total_cost_value = self.format_sum(total_cost_value_raw)

        comparative_value_raw = getattr(self, "comparative_value", 0)
        comparative_value = self.format_sum(comparative_value_raw)
        if method == "Среднее арифметическое":
            cost_percent = 50
            comparative_percent = 50
            weighted_cost = self.cost_value / 2 if self.checkBox_cost.isChecked() else 0
            weighted_comparative = self.comparative_value / 2 if self.checkBox_comparative.isChecked() else 0
        else:
            cost_percent = self.spinBox_cost_percent.value()
            comparative_percent = self.spinBox_comparative_percent.value()
            weighted_cost = self.weighted_cost
            weighted_comparative = self.weighted_comparative



        return {
            "method": self.comboBox_agreement_list.currentText(),
            "use_cost": self.checkBox_cost.isChecked(),
            "use_comparative": self.checkBox_comparative.isChecked(),
            "cost_percent": cost_percent,
            "comparative_percent": comparative_percent,            
            "final_cost": final_cost_text,
            "edited_final_cost": edited_cost,
            "amount_in_words": words,
            "building_cost": round(building_cost_raw),
            "land_cost": round(land_cost_raw),
            "total_cost_value": round(total_cost_value_raw),
            "comparative_final_cost_value": round(comparative_value_raw),
            "agreement_method_summary": agreement_method_summary,
            "weighted_cost": round(weighted_cost),
            "weighted_comparative": round(weighted_comparative),
            "weights": {
            "cost": cost_percent,
            "comparative": comparative_percent
                }
        }




    def load_agreement_data(self, data):
        self._loading = True
        try:
            # ОТКЛЮЧАЕМ СИГНАЛЫ
            self.checkBox_cost.blockSignals(True)
            self.checkBox_comparative.blockSignals(True)
            self.comboBox_agreement_list.blockSignals(True)

            self.spinBox_cost_percent.blockSignals(True)
            self.spinBox_comparative_percent.blockSignals(True)

            # Загружаем значения
            self.comboBox_agreement_list.setCurrentText(data.get("method", "Взвешенное среднее"))
            self.checkBox_cost.setChecked(data.get("use_cost", True))
            self.checkBox_comparative.setChecked(data.get("use_comparative", True))

            weights = data.get("weights", {})
            cost_weight = weights.get("cost", data.get("cost_percent", 50))
            comparative_weight = weights.get("comparative", data.get("comparative_percent", 50))

            self.spinBox_cost_percent.setValue(cost_weight)
            self.spinBox_comparative_percent.setValue(comparative_weight)

            self.cost_value = float(data.get("total_cost_value", 0))
            self.comparative_value = float(data.get("comparative_final_cost_value", 0))

            self.checkBox_cost.setText(f"({self.format_sum(self.cost_value)} сум)\n\n")
            self.checkBox_comparative.setText(f"{self.format_sum(self.comparative_value)} сум")

            # Восстановление building_cost и land_cost
            self.building_cost = float(data.get("building_cost", 0))
            self.land_cost = float(data.get("land_cost", 0))
            self.total_cost_value = float(data.get("total_cost_value", 0))
            self.label_building_land.setText(
                f"Стоимость улучшений: {self.format_sum(self.building_cost)} сум\n"
                f"Права на землю: {self.format_sum(self.land_cost)} сум"
            )

            # Восстановление финального текста стоимости
            self.label_final_cost.setText(data.get("final_cost", "Итоговая стоимость: н/д"))

            # Восстановление взвешенных стоимостей
            self.weighted_cost = float(data.get("weighted_cost", 0))
            self.weighted_comparative = float(data.get("weighted_comparative", 0))

            # Можно сохранить также amount_in_words и summary, если они нужны
            self.amount_in_words = data.get("amount_in_words", "")
            self.agreement_method_summary = data.get("agreement_method_summary", "")

            if sum([self.checkBox_cost.isChecked(), self.checkBox_comparative.isChecked()]) > 1:
                self.recalculate_weighted_average()
            else:
                self.label_final_cost.setText(data.get("final_cost", "Итоговая стоимость: н/д"))

        except Exception as e:
            print(f"[ERROR] Не удалось загрузить данные вкладки Согласование: {e}")
        finally:
            # ВКЛЮЧАЕМ ОБРАТНО СИГНАЛЫ
            self.spinBox_cost_percent.blockSignals(False)
            self.spinBox_comparative_percent.blockSignals(False)
            self.checkBox_cost.blockSignals(False)
            self.checkBox_comparative.blockSignals(False)
            self.comboBox_agreement_list.blockSignals(False)

            self._loading = False





    def final_save(self):
        self.valuation_window.save_report()
        self.valuation_window.main_window.update_last_valuation_cost_from_agreement()



# ФОРМИРУЕМ ОТЧЁТ

    

    def insert_kadastr_table_into_word(self, docx_path, html_table, marker="{{TABLE_KADASTR}}"):
        from bs4 import BeautifulSoup
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        soup = BeautifulSoup(html_table, "html.parser")
        table_html = soup.find("table")
        if not table_html:
            return

        rows = table_html.find_all("tr")
        if not rows:
            return

        doc = Document(docx_path)

        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                p = paragraph._element
                p_parent = p.getparent()

                # Добавляем таблицу в XML после маркера
                table = doc.add_table(rows=0, cols=2)
                tbl_element = table._element
                p.addnext(tbl_element)

                # Оформление таблицы
                try:
                    table.style = 'Table Grid'  # работает, если есть в шаблоне
                except:
                    pass

                table.alignment = WD_TABLE_ALIGNMENT.CENTER  # центрируем таблицу

                for tr in rows:
                    cols = tr.find_all("td")
                    if len(cols) != 2:
                        continue
                    row_cells = table.add_row().cells
                    for i in range(2):
                        text = cols[i].get_text(strip=True)
                        row_cells[i].text = text
                        # Центрируем текст внутри ячеек
                        for paragraph in row_cells[i].paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        row_cells[i].width = Inches(3.0)

                # Удаляем маркер
                p_parent.remove(p)
                break

        doc.save(docx_path)


    

    def load_market_analysis(self, oblast_name, research_folder="research"):
        """Ищет и загружает содержимое анализа по названию области"""
        import os
        from docx import Document

        docx_filename = f"{oblast_name}.docx"

        # Получаем путь к research, находящемуся рядом с .exe или main.py
        base_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
        docx_path = os.path.join(base_dir, research_folder, docx_filename)

        if not os.path.exists(docx_path):
            return f"⚠️ Анализ по области «{oblast_name}» не найден."

        doc = Document(docx_path)
        full_text = [para.text for para in doc.paragraphs]
        
        return "\n\n".join(full_text)


  
    def insert_comparative_table(self, docx_path, comparative, marker="[ comparative_table ]"):
        doc = Document(docx_path)

        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                parent.remove(paragraph._element)

                vertical_headers = comparative.get("vertical_headers", [])
                horizontal_headers = comparative.get("horizontal_headers", [])
                table_data = comparative.get("table_data", [])
                final_value = comparative.get("label_comparative_final_cost", "—")

                if not vertical_headers or not horizontal_headers or not table_data:
                    continue

                cols = len(horizontal_headers) + 1  # +1 для названий строк
                rows = len(vertical_headers)

                table = doc.add_table(rows=0, cols=cols)
                table.style = "Table Grid"
                table.autofit = True

                # Первая строка — заголовки
                header_row = table.add_row().cells
                header_row[0].text = "Параметры"
                for i, analog in enumerate(horizontal_headers):
                    header_cell = header_row[i + 1]
                    header_cell.text = analog["text"]
                    if analog.get("url"):
                        run = header_cell.paragraphs[0].runs[0]
                        run.font.underline = True
                        run.font.color.rgb = None  # можно задать цвет, если нужно

                # Данные таблицы
                for row_index, row_label in enumerate(vertical_headers[:-1]):
                    row_cells = table.add_row().cells
                    row_cells[0].text = row_label
                    for col_index in range(len(horizontal_headers)):
                        value = table_data[row_index][col_index]
                        row_cells[col_index + 1].text = value

                # Последняя строка — объединённая
                final_row = table.add_row().cells
                merged_cell = final_row[0].merge(final_row[-1])
                paragraph = merged_cell.paragraphs[0]
                run = paragraph.add_run(final_value)
                run.bold = True
                run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Вставка таблицы на место маркера
                parent.insert(index, table._element)
                break

        doc.save(docx_path)


    def improved_insert_land_table(self, docx_path, land_valuation, marker="[ LAND_TABLE ]"):
        doc = Document(docx_path)

        # Найти маркер
        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                parent.remove(paragraph._element)

                analogs = land_valuation["horizontal_headers"]
                vertical_headers = land_valuation["vertical_headers"]
                table_data = land_valuation["table_data"]

                # Создание таблицы: параметры + каждый аналог = n+1 столбцов
                table = doc.add_table(rows=0, cols=len(analogs) + 1)
                table.style = "Table Grid"
                table.autofit = True

                # Первая строка — заголовки
                header_row = table.add_row().cells
                header_row[0].text = "Параметры"
                for i, analog in enumerate(analogs):
                    header_row[i + 1].text = analog["text"]

                # Основные строки таблицы
                for row_index, row_label in enumerate(vertical_headers[:-1]):  # исключаем последнюю строку
                    row_cells = table.add_row().cells
                    row_cells[0].text = row_label
                    for col_index in range(len(analogs)):
                        value = table_data[row_index][col_index]
                        row_cells[col_index + 1].text = value

                # Последняя строка — объединённая ячейка со средней стоимостью
                avg_row = table.add_row().cells
                avg_label_cell = avg_row[0]
                avg_label_cell.merge(avg_row[-1])
                avg_paragraph = avg_label_cell.paragraphs[0]
                run = avg_paragraph.add_run(
                    "Средняя стоимость за 1 сотку: " +
                    land_valuation.get("cost_per_sotka_soum", "—")
                )

                run.bold = True
                avg_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Вставка таблицы
                parent.insert(index, table._element)
                break

        doc.save(docx_path)


     

    def insert_agreement_table(self, docx_path, agreement_data, marker='[ agreement_table ]'):
        doc = Document(docx_path)

        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                parent.remove(paragraph._element)

                use_cost = agreement_data.get("use_cost", False)
                use_comparative = agreement_data.get("use_comparative", False)

                if use_cost and use_comparative:
                    # Расширенная таблица
                    table = doc.add_table(rows=2, cols=7)
                    table.style = "Table Grid"
                    table.autofit = True

                    headers = [
                        "Затратный подход", "", "",
                        "Сравнительный подход", "", "",
                        "Согласованная стоимость"
                    ]
                    subheaders = [
                        "стоимость улучшений + стоимость права на земельный участок", "Удельный вес", "Взвешенное значение",
                        "Стоимость объекта", "Удельный вес", "Взвешенное значение",
                        ""
                    ]
                    for i, text in enumerate(headers):

                        cell = table.cell(0, i)
                        cell.text = text
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for i, text in enumerate(subheaders):

                        cell = table.cell(1, i)
                        cell.text = text
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Данные
                    row = table.add_row().cells
                    def format_sum(value):
                        try:
                            return f"{round(float(str(value).replace(' ', ''))):,}".replace(",", " ")
                        except:
                            return str(value)

                    row[0].text = format_sum(agreement_data.get("total_cost_value", 0))
                    row[1].text = str(agreement_data.get("cost_percent", ""))
                    row[2].text = format_sum(agreement_data.get("weighted_cost", 0))
                    row[3].text = format_sum(agreement_data.get("comparative_final_cost_value", 0))
                    row[4].text = str(agreement_data.get("comparative_percent", ""))
                    row[5].text = format_sum(agreement_data.get("weighted_comparative", 0))
                    row[6].text = format_sum(agreement_data.get("edited_final_cost", 0))


                else:
                    # Упрощённая таблица
                    table = doc.add_table(rows=3, cols=3)
                    table.style = "Table Grid"
                    table.autofit = True

                    headers = ["Объект оценки", "", "Стоимость согласования"]
                    subheaders = ["", "Стоимость", ""]

                    for i, text in enumerate(headers):
                        

                        table.cell(0, i).text = text
                        table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    for i, text in enumerate(subheaders):
                        

                        table.cell(1, i).text = text
                        table.cell(1, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    method = "Затратный подход" if use_cost else "Сравнительный подход"
                    cost = agreement_data.get("edited_final_cost", "")
                    row = table.rows[2].cells
                    row[0].text = "Жилой дом"
                    row[1].text = cost
                    row[2].text = cost

                parent.insert(index, table._element)
                break

        doc.save(docx_path)



        

    def insert_koeff_table(self, docx_path, koefs_data, marker="[ koeff_table ]"):
        doc = Document(docx_path)

        # Проверка наличия данных
        koefs_table = koefs_data.get("koefs_table", [])
        if not koefs_table:
            return

        for paragraph in doc.paragraphs:
            if marker in paragraph.text:
                parent = paragraph._element.getparent()
                index = parent.index(paragraph._element)
                parent.remove(paragraph._element)

                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                table.autofit = True

                # Заголовки
                header = table.rows[0].cells
                header[0].text = "Дата"
                header[1].text = "Коэффициент"

                # Данные
                for row in koefs_table:
                    if len(row) != 2:
                        continue
                    date, value = row
                    cells = table.add_row().cells
                    cells[0].text = date
                    cells[1].text = value

                # Вставка таблицы в документ
                parent.insert(index, table._element)
                break

        doc.save(docx_path)

    

    def on_upload_report_clicked(self):
        self.valuation_window.save_report()  # Просто вызываем уже существующий, не тронутый метод

        agreement = self.collect_agreement_data()
        use_cost = agreement.get("use_cost", False)
        use_comparative = agreement.get("use_comparative", False)

        if not use_cost or not use_comparative:
            dialog = MethodRejectionDialog(self, valuation_window=self.valuation_window)
            if dialog.exec_() == QDialog.Accepted:
                self.selected_template = dialog.selected_template

            else:
                QMessageBox.warning(self, "Отмена", "Выгрузка отменена.")
                return
        
                # Загрузка отчётных данных
        valuation_window = self.valuation_window
        report_number = valuation_window.report_number_input.text().strip()
        report_path = valuation_window.main_window.report_manager.get_report_path(report_number)

        if not os.path.exists(report_path):
            QMessageBox.warning(self, "Ошибка", f"Файл отчёта не найден:\n{report_path}")
            return

        with open(report_path, "r", encoding="utf-8") as f:
            report_data = json.load(f)

        agreement = report_data.get("agreement", {})
        use_cost = agreement.get("use_cost", True)
        use_comparative = agreement.get("use_comparative", True)
               
       
        template_name = getattr(self, "selected_template", "result.docx")
        template_path = os.path.join("reports", template_name)
        
        if not os.path.exists(template_path):
            QMessageBox.critical(self, "Ошибка", f"Файл шаблона {template_name} не найден.")
            return
        if not is_license_valid():
            QMessageBox.warning(
                self,
                "Подписка истекла",
                "Срок действия подписки истёк или не подтверждён.\nПожалуйста, оплатите подписку."
            )

            dialog = PaymaentDialog(self)
            result = dialog.exec_()

            # Если пользователь не завершил оплату — выходим
            # if result != QDialog.Accepted or not getattr(dialog, "paid", False):
            #     QMessageBox.warning(self, "Оплата", "Оплата не завершена. Выгрузка отчёта невозможна.")
            #     return  # 👈 Блокируем выгрузку

            # 👇 Здесь может быть проверка повторной валидации или запрос на сервер
            if not is_license_valid():
                # QMessageBox.critical(self, "Ошибка", "Оплата ещё не подтверждена. Попробуйте позже.")
                return

                  
        def generate_engineering_description(comms: dict, heating: str) -> str:
            mapping = {
                "газификация": "Газоснабжение",
                "электроосвещение": "Электроосвещение",
                "водоснабжение": "Водоснабжение",
                "канализация": "Канализация",
                "телефонная_линия": "Телефонная линия",
                "электрический_водонагреватель": "Электрический водонагреватель",
                "горячее_водоснабжение": "Горячее водоснабжение"
            }

            available = [name for key, name in mapping.items() if comms.get(key)]
            count = len(available)

            if count <= 2:
                level = "низкая"
            elif count <= 4:
                level = "удовлетворительная"
            else:
                level = "хорошая"

            lines = ["Инженерная инфраструктура:"]
            if available:
                lines.append("Объект обеспечен следующими инженерными коммуникациями:")
                lines.extend(f"– {item}" for item in available)
            else:
                lines.append("Объект не обеспечен инженерными коммуникациями.")

            lines.append("")
            lines.append(f"Отопление: {heating or 'не указано'}")
            lines.append(f"Степень обеспеченности инженерной инфраструктурой — {level}.")

            return "\n".join(lines)
        
        


       
       

        def insert_liter_tables(docx_path, liters):
            doc = Document(docx_path)

            def remove_last_paragraph_if_empty(doc):
                if doc.paragraphs and not doc.paragraphs[-1].text.strip():
                    p = doc.paragraphs[-1]._element
                    p.getparent().remove(p)

            def insert_full_liter(doc, parent, index, liter):
                def add_row(table, key, value):
                    row = table.add_row().cells
                    row[0].text = key
                    row[1].text = str(value)

                def format2(val):
                    return f"{float(val):.2f}" if val not in ("", None) else ""

                # --- Таблица характеристик ---
                table = doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                table.autofit = True

                filters = liter.get("filters", {})
                m = liter.get("measurements", {})

                add_row(table, "Тип строения", liter["building_type"])
                if filters.get("Этажность"):
                    add_row(table, "Этажность", filters["Этажность"])
                if filters.get("Отделка"):
                    add_row(table, "Отделка", filters["Отделка"])
                if filters.get("Толщина стен"):
                    add_row(table, "Толщина стен", filters["Толщина стен"])
                if filters.get("Материал стен"):
                    add_row(table, "Материал стен", filters["Материал стен"])
                if filters.get("Примыкание"):
                    add_row(table, "Примыкание", filters["Примыкание"])
                if filters.get("Фундамент"):
                    add_row(table, "Фундамент", filters["Фундамент"])
                if filters.get("Кровля"):
                    add_row(table, "Кровля", filters["Кровля"])
                if filters.get("Перекрытие"):
                    add_row(table, "Перекрытие", filters["Перекрытие"])
                if filters.get("Тип покрытия"):
                    add_row(table, "Тип покрытия", filters["Тип покрытия"])
                if filters.get("Высота стен"):
                    add_row(table, "Высота стен", filters["Высота стен"])
                if m.get("square"):
                    add_row(table, "Площадь", m["square"])
                if m.get("length"):
                    add_row(table, "Протяжённость", m["length"])
                if m.get("height"):
                    add_row(table, "Высота", m["height"])
                if m.get("volume"):
                    add_row(table, "Объём", m["volume"])
                if m.get("ukup_price_label"):
                    add_row(table, "Стоимость по УКУП", m["ukup_price_label"])

                remove_last_paragraph_if_empty(doc)
                parent.insert(index, table._element)
                index += 1

                # --- Описание аналога ---
                html = liter.get("analog_description_html", "")
                if html:
                    soup = BeautifulSoup(html, "html.parser")
                    desc_title = doc.add_paragraph("")
                    parent.insert(index, desc_title._element)
                    index += 1

                    paragraphs = soup.get_text(separator="\n").split("\n")
                    for text in paragraphs:
                        if text.strip():
                            para = doc.add_paragraph(text.strip())
                            parent.insert(index, para._element)
                            index += 1

                # --- Таблица корректировок ---
                analog_table = doc.add_table(rows=1, cols=3)
                analog_table.style = "Table Grid"
                analog_table.autofit = True

                header = analog_table.rows[0].cells
                header[0].text = "Наименование"
                header[1].text = "Значение"
                header[2].text = "Комментарий"

                def add_analog_row(name, value, comment):
                    row = analog_table.add_row().cells
                    row[0].text = name
                    row[1].text = str(value)
                    row[2].text = comment

                facade_price = liter.get("facade_corrected_price", 0)
                if facade_price:
                    add_analog_row(f"Корректировка на фасад", format2(facade_price), liter.get("facade_type", ""))

                # --- Детализированные инженерные улучшения ---
               
                improvements = liter.get("improvement_details", [])
                if improvements:
                    for imp in improvements:
                        name = imp.get("name", "Неизвестное улучшение")
                        percent = imp.get("correction_percent", 0.0)
                        absolute = imp.get("correction_value", 0.0)
                        add_analog_row(name, f"{absolute:,.2f} сум", f"{percent * 100:+.2f}%")


                # --- Детализированные отклонения ---
                deviations = liter.get("deviation_details", [])
                if deviations:
                    for dev in deviations:
                        if dev.get("selected", False):  # только выбранные
                            name = dev.get("name", "Отклонение")
                            value = dev.get("value", 0)
                            add_analog_row(name, f"{value:+,.2f} сум", "Отклонение")

                height_correction = liter.get("height_corrected_price", 0)
                if height_correction:
                    add_analog_row("Корректировка по высоте", format2(height_correction), "сум")

                add_analog_row("Откорректированная стоимость УКУП", format2(liter["corrected_price"]), "сум")
                add_analog_row(liter["reg_coeff_type"], liter["reg_coeff"], "")
                add_analog_row(liter["stat_koeff_label"], f"{liter['stat_coeff']:.3f}", "")
                add_analog_row("Прибыль предпринимателя", f"{liter['developer_percent']}", "%")
                add_analog_row("Коэффициент сейсмичности", f"{liter['sesmos']}", "")
                add_analog_row("Восттановительная стоимость", f"{self.format_sum(liter['replacement_cost'])}", "Сум")

                remove_last_paragraph_if_empty(doc)
                parent.insert(index, analog_table._element)
                index += 1

                # --- Подзаголовок ---
                para3 = doc.add_paragraph("РАСЧЁТ ФИЗИЧЕСКОГО ИЗНОСА", style="Heading 3")
                parent.insert(index, para3._element)
                index += 1

                # --- Таблица износа ---
                se_table = doc.add_table(rows=1, cols=4)
                se_table.style = "Table Grid"
                se_table.autofit = True

                se_header = se_table.rows[0].cells
                se_header[0].text = "Конструктивные элементы"
                se_header[1].text = "Доля %"
                se_header[2].text = "Поправка"
                se_header[3].text = "Износ %"

                for se in liter["structural_elements"]:
                    row = se_table.add_row().cells
                    row[0].text = se["Конструкции"]
                    row[1].text = str(se["Доля %"])
                    row[2].text = str(
                        se.get("Поправка к удельным весам %") or
                        se.get("Поправка к\nудельным весам %") or
                        "—"
                    )
                    row[3].text = str(
                        se.get("Физический износ %") or
                        se.get("Физический\nизнос %") or
                        "—"
                    )


                remove_last_paragraph_if_empty(doc)
                parent.insert(index, se_table._element)
                index += 1

                # --- Таблица итогов ---
                summary_table = doc.add_table(rows=4, cols=2)
                summary_table.style = "Table Grid"
                summary_table.autofit = True

                summary_table.rows[0].cells[0].text = "Коэф. недостроенности"
                summary_table.rows[0].cells[1].text = str(liter.get("inconsistency", ""))

                summary_table.rows[1].cells[0].text = "Общий процент износа"
                summary_table.rows[1].cells[1].text = f"{liter['wear_percent']:,.2f} %"

                summary_table.rows[2].cells[0].text = "Физический износ"
                summary_table.rows[2].cells[1].text = f"{liter['wear_price']:,.2f} сум"

                summary_table.rows[3].cells[0].text = "Оценочная стоимость"
                cell = summary_table.rows[3].cells[1]
                cell.text = ""
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(f"{liter['final_cost']:,.2f} сум")
                run.bold = True

                remove_last_paragraph_if_empty(doc)
                parent.insert(index, summary_table._element)
                index += 1

                # --- Финальный отступ (один параграф) ---
                empty = doc.add_paragraph()
                parent.insert(index, empty._element)
                index += 1
                return index

            # --- Вставка на место маркера ---
            for paragraph in doc.paragraphs:
                if "[[LITER_TABLES_PLACEHOLDER]]" in paragraph.text:
                    parent = paragraph._element.getparent()
                    index = parent.index(paragraph._element)
                    parent.remove(paragraph._element)

                    for liter in liters:
                        para = doc.add_paragraph(f"Литер №{liter['number']}", style='Heading 2')
                        parent.insert(index, para._element)
                        index += 1
                        index = insert_full_liter(doc, parent, index, liter)  # <<< ОБНОВЛЯЕМ INDEX


                    break

            

            doc.save(docx_path)


        try:
            # 1. Получаем номер отчёта и регномер из ValuationMainWindow
            valuation_window = self.valuation_window
            report_number = valuation_window.report_number_input.text().strip()
            reg_number = valuation_window.lineEdit_reg_number.text().strip() or f"Report_{report_number}"

            # 2. Загружаем данные из report_{номер}.json
            report_path = valuation_window.main_window.report_manager.get_report_path(report_number)
            if not os.path.exists(report_path):
                QMessageBox.warning(self, "Ошибка", f"Файл отчёта не найден:\n{report_path}")
                return

            with open(report_path, "r", encoding="utf-8") as f:
                report_data = json.load(f)

            # 3. Проверяем наличие шаблона result.docx
            # template_path = os.path.join("reports", "result.docx")
            if not os.path.exists(template_path):
                QMessageBox.critical(self, "Ошибка", "Файл шаблона result.docx не найден.")
                return

            # 4. Папка сохранения
            save_dir = valuation_window.main_window.save_directory
            if not save_dir:
                QMessageBox.critical(self, "Ошибка", "Путь сохранения не задан.")
                return
            reg_number = valuation_window.lineEdit_reg_number.text().strip() or f"Report_{report_number}"
            report_folder = os.path.join(save_dir, reg_number)
            os.makedirs(report_folder, exist_ok=True)
            output_path = os.path.join(report_folder, f"Отчёт №{report_number}.docx")

            # 5. Подготовка контекста
            
            liters = report_data.get("liters", [])
            lines = []

            for liter in liters:
                number = liter.get("number", "")
                building_type = liter.get("building_type", "")
                measurements = liter.get("measurements", {})
                square = measurements.get("square", "")
                length = measurements.get("length", "")

                if building_type.lower().startswith("огражд") or building_type.lower() == "забор":
                    value = length or "—"
                    unit = "м"
                    label = "протяжённость"
                else:
                    value = square or "—"
                    unit = "м²"
                    label = "площадь"

                line = f"№{number}. {building_type}, {label}: {value} {unit}"
                lines.append(line)

            liters_block_text = "\n".join(lines)
            def sanitize_html_table(raw_html):
                # Удалим \n и лишние пробелы между тегами
                cleaned = raw_html.replace('\n', ' ').replace('\r', ' ')
                cleaned = ' '.join(cleaned.split())  # удалим множественные пробелы

                # Заменим проблемные символы
                cleaned = cleaned.replace("‘", "'").replace("’", "'").replace("`", "'")
                cleaned = cleaned.replace("“", '"').replace("”", '"')
                cleaned = cleaned.replace("*", "")  # или удалить/заменить как нужно

                # Escape HTML спецсимволы внутри значений, если это нужно
                return cleaned
            raw_table_html = sanitize_html_table(report_data.get("kadastr_table_html", ""))
            rendered_table_html = Template("{{ val | safe }}").render(val=raw_table_html)
            oblast_name = report_data.get("administrative", {}).get("oblast", "")
            regional_market_analysis = self.load_market_analysis(oblast_name)

            communications = report_data.get("communications", {})
            heating = report_data.get("heating", "")
            engineering_description = generate_engineering_description(communications, heating)
            rayon = report_data.get("administrative", {}).get("rayon", "").lower()
            if "город" in rayon:
                density = "высокая"
            else:
                density = "средняя"
            agreement = report_data.get("agreement", {})
            use_cost = agreement.get("use_cost", False)
            use_comparative = agreement.get("use_comparative", False)

            # Если отключён хотя бы один подход, спрашиваем причины
            


            context = {
                "liters_block": liters_block_text,
                "liters": report_data.get("liters", []),
                "report_number": report_number,
                "reg_number": reg_number,
                "contract_date": report_data.get("contract_date", ""),
                "inspection_date": report_data.get("inspection_date", ""),
                "exchange_rate": report_data.get("exchange_rate", ""),
                "address": report_data.get("address", ""),
                "owner_name": report_data.get("owner_name", ""),
                "valuation_purpose": report_data.get("valuation_purpose", ""),
                "price_type": report_data.get("price_type", ""),
                "buyer_type": report_data.get("buyer_type", ""),
                "buyer_name": report_data.get("buyer_name", ""),
                "buyer_passport_series": report_data.get("buyer_passport_series", ""),
                "buyer_passport_number": report_data.get("buyer_passport_number", ""),
                "buyer_address": report_data.get("buyer_address", ""),
                "land_area": report_data.get("land_area", ""),
                "total_area": report_data.get("total_area", ""),
                "useful_area": report_data.get("useful_area", ""),
                "living_area": report_data.get("living_area", ""),
                "cadastral_number": report_data.get("cadastral_number", ""),
                "profit": report_data.get("profit", ""),
                 "agreement": {
                    "method": report_data.get("agreement", {}).get("method", ""),
                    "cost_percent": report_data.get("agreement", {}).get("cost_percent", ""),
                    "comparative_percent": report_data.get("agreement", {}).get("comparative_percent", ""),
                    "edited_final_cost": report_data.get("agreement", {}).get("edited_final_cost", ""),
                    "amount_in_words": report_data.get("agreement", {}).get("amount_in_words", ""),
                    "building_cost": (
                        self.format_sum(agreement.get("building_cost") or 0)
                        if use_cost else "Затратный подход не использовался"
                    ),
                    "land_cost": (
                        self.format_sum(agreement.get("land_cost", ""))
                        if use_cost else "Затратный подход не использовался"
                    ),
                    "total_cost_value": (
                        self.format_sum(agreement.get("total_cost_value", ""))
                        if use_cost else "Затратный подход не использовался"
                    ),
                    "comparative_final_cost_value": (
                        self.format_sum(agreement.get("comparative_final_cost_value", ""))
                        if use_comparative else "Сравнительный подход не использовался"
                    ),


                },
                'analogs_count': report_data.get("land_valuation", {}).get('analogs_count', ""),

                "administrative": {
                        "oblast": report_data.get("administrative", {}).get("oblast", ""),
                        "rayon": report_data.get("administrative", {}).get("rayon", "")
                    }, 
                "regional_market_analysis": regional_market_analysis,
                "engineering_description": engineering_description,
                "density": density,
                "lineEdit_CBUF": report_data.get("lineEdit_CBUF", ""),
                "profit": report_data.get("profit"),
                


            }
            

            
            


            context["TABLE_KADASTR"] = "{{ TABLE_KADASTR }}"  # чтобы оставить маркер в документе
            context["agreement_method_summary"] = report_data["agreement"].get("agreement_method_summary", "")

            # 6. Рендеринг
            env = Environment(undefined=DebugUndefined)
            doc = DocxTemplate(template_path)
            doc.env = env
            doc.render(context)
            
            doc.save(output_path)
            insert_liter_tables(output_path, report_data["liters"])
           

            self.improved_insert_land_table(output_path, report_data["land_valuation"], marker="[ LAND_TABLE ]")
            self.insert_koeff_table(output_path, report_data["koefs"], marker="[ koeff_table ]")
            self.insert_comparative_table(output_path, report_data['comparative'], marker="[ comparative_table ]")
            self.insert_kadastr_table_into_word(output_path, raw_table_html, marker="{{ TABLE_KADASTR }}")
            self.insert_agreement_table(output_path, report_data["agreement"], marker='[ agreement_table ]')
            QMessageBox.information(self, "Успех", f"Отчёт успешно создан:\n{output_path}")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сгенерировать отчёт:\n{str(e)}")



    